"""Build .docx and .pdf artifacts from a translated body of text.

The translation pass (see anna_chat/translate.py) asks Bedrock to mark
basic structure in the output using lightweight Markdown:

  # heading                  → Heading 1 (docx) / large bold (pdf)
  ## subheading              → Heading 2 / medium bold
  - bullet item              → bulleted list
  1. numbered item           → numbered list
  **bold** / *italic*        → inline bold / italic runs
  paragraph (anything else)  → body paragraph

This module parses that subset and renders into real Word/PDF styling so
the output documents look structured rather than as one wall of body text.

We do NOT pull in a full Markdown library (e.g. python-markdown or
mistune): the surface area is tightly bounded, every supported
construct is one-line-deterministic, and a 60-line hand-rolled parser
keeps the Lambda zip small and the code readable.

RTL handling: when the target language is in `RTL_LANGUAGE_CODES` we set
`<w:bidi/>` on docx paragraphs (python-docx 1.x doesn't expose
`paragraph_format.rtl`) and right-align the pdf body. Full bidi shaping
in PDF is out of scope for V1 — see TRANSLATE_CONTRACT.md known
limitations.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph as DocxParagraph
from lxml import etree
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    TableStyle,
)
from reportlab.platypus import (
    Table as PdfTable,
)

# ISO-639 codes whose scripts run right-to-left. Kept narrow on purpose.
RTL_LANGUAGE_CODES: frozenset[str] = frozenset({"ar", "he"})

_FOOTER_LINE = "Translated by Praxis"

# Inline markdown matchers. Bold MUST be tested before italic — `**foo**`
# would otherwise be parsed as `*` `*foo*` `*`.
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)")
_NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*•]\s+(.+)$")


# ──────────────────────────────────────────────────────────────────────────
# Block parsing — turns the markdown body into a flat list of typed blocks
# ──────────────────────────────────────────────────────────────────────────


BlockKind = Literal["h1", "h2", "bullet", "numbered", "paragraph", "table"]


@dataclass(frozen=True)
class Block:
    kind: BlockKind
    text: str = ""
    # Tables carry rows here instead of in `text`. Each row is a list of
    # cell strings (raw inline-markdown; bold/italic NOT yet expanded).
    # Convention: rows[0] is the header, rows[1:] are body rows.
    rows: tuple[tuple[str, ...], ...] = ()


_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


def _split_table_row(line: str) -> list[str]:
    """Split a markdown table row `| a | b\\| c |` into cells, honoring `\\|`.

    Walks the string char by char so an escaped `\\|` inside a cell isn't
    mistaken for a column boundary. Strips outer pipes and per-cell
    whitespace; leaves backslash escapes resolved (`\\|` → `|`, `\\\\` → `\\`).
    """
    # Strip leading/trailing whitespace and the bordering pipes.
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    cells: list[str] = []
    buf: list[str] = []
    i = 0
    n = len(s)
    while i < n:
        c = s[i]
        if c == "\\" and i + 1 < n:
            buf.append(s[i + 1])
            i += 2
            continue
        if c == "|":
            cells.append("".join(buf).strip())
            buf = []
            i += 1
            continue
        buf.append(c)
        i += 1
    cells.append("".join(buf).strip())
    return cells


def _parse_blocks(body: str) -> list[Block]:
    """Parse the markdown body into a flat list of typed blocks.

    Recognized constructs (one block per match): h1 / h2 / bullet /
    numbered / table / paragraph. Tables require the standard
    GitHub-flavored two-line opener (header row + separator row) — a
    line that LOOKS like a table row but isn't preceded by a separator
    is treated as a paragraph, which is the right call: a stray pipe
    in body text shouldn't fold into a phantom 1-row table.

    Adjacent plain lines are coalesced into one paragraph. Blank lines
    flush whatever paragraph was being accumulated.
    """
    if not body:
        return []
    norm = body.replace("\r\n", "\n").replace("\r", "\n")
    lines = norm.split("\n")
    blocks: list[Block] = []
    para_lines: list[str] = []

    def flush_paragraph() -> None:
        if not para_lines:
            return
        text = " ".join(line.strip() for line in para_lines).strip()
        if text:
            blocks.append(Block("paragraph", text))
        para_lines.clear()

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        # Table: header row + separator row + zero or more body rows.
        if (
            _TABLE_ROW_RE.match(stripped)
            and i + 1 < n
            and _TABLE_SEP_RE.match(lines[i + 1].strip())
        ):
            flush_paragraph()
            header = tuple(_split_table_row(stripped))
            j = i + 2
            body_rows: list[tuple[str, ...]] = []
            while j < n and _TABLE_ROW_RE.match(lines[j].strip()):
                body_rows.append(tuple(_split_table_row(lines[j].strip())))
                j += 1
            blocks.append(
                Block("table", rows=(header, *tuple(body_rows)))
            )
            i = j
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            blocks.append(Block("h2", stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            blocks.append(Block("h1", stripped[2:].strip()))
            i += 1
            continue
        bullet = _BULLET_RE.match(stripped)
        if bullet:
            flush_paragraph()
            blocks.append(Block("bullet", bullet.group(1).strip()))
            i += 1
            continue
        numbered = _NUMBERED_RE.match(stripped)
        if numbered:
            flush_paragraph()
            blocks.append(Block("numbered", numbered.group(2).strip()))
            i += 1
            continue
        para_lines.append(stripped)
        i += 1

    flush_paragraph()
    return blocks


# ──────────────────────────────────────────────────────────────────────────
# Inline markdown → typed runs (used by the docx renderer)
# ──────────────────────────────────────────────────────────────────────────


@dataclass(frozen=True)
class InlineRun:
    text: str
    bold: bool = False
    italic: bool = False


def _parse_inline(text: str) -> list[InlineRun]:
    """Split text into runs, marking bold and italic spans.

    Pass 1: replace `**...**` with placeholders carrying a "bold" flag.
    Pass 2: replace `*...*` likewise. Then walk the resulting flat
    structure and emit InlineRun records. Order matters because the
    bold pattern (two stars) must match before italic (one star).
    """
    if not text:
        return []

    # We use a sentinel marker the source text won't contain.
    # \x00\x01 are control chars stripped at the parser boundary.
    SENTINEL_OPEN_BOLD = "\x00B\x01"
    SENTINEL_CLOSE_BOLD = "\x00b\x01"
    SENTINEL_OPEN_ITALIC = "\x00I\x01"
    SENTINEL_CLOSE_ITALIC = "\x00i\x01"

    def _repl_bold(m: re.Match[str]) -> str:
        inner = m.group(1)
        return f"{SENTINEL_OPEN_BOLD}{inner}{SENTINEL_CLOSE_BOLD}"

    def _repl_italic(m: re.Match[str]) -> str:
        inner = m.group(1)
        return f"{SENTINEL_OPEN_ITALIC}{inner}{SENTINEL_CLOSE_ITALIC}"

    s = _BOLD_RE.sub(_repl_bold, text)
    s = _ITALIC_RE.sub(_repl_italic, s)

    runs: list[InlineRun] = []
    bold = False
    italic = False
    buf: list[str] = []

    def flush() -> None:
        chunk = "".join(buf)
        buf.clear()
        if chunk:
            runs.append(InlineRun(chunk, bold=bold, italic=italic))

    i = 0
    n = len(s)
    while i < n:
        # Look for a sentinel at position i
        if s.startswith(SENTINEL_OPEN_BOLD, i):
            flush()
            bold = True
            i += len(SENTINEL_OPEN_BOLD)
            continue
        if s.startswith(SENTINEL_CLOSE_BOLD, i):
            flush()
            bold = False
            i += len(SENTINEL_CLOSE_BOLD)
            continue
        if s.startswith(SENTINEL_OPEN_ITALIC, i):
            flush()
            italic = True
            i += len(SENTINEL_OPEN_ITALIC)
            continue
        if s.startswith(SENTINEL_CLOSE_ITALIC, i):
            flush()
            italic = False
            i += len(SENTINEL_CLOSE_ITALIC)
            continue
        buf.append(s[i])
        i += 1
    flush()
    return runs


# ──────────────────────────────────────────────────────────────────────────
# DOCX rendering
# ──────────────────────────────────────────────────────────────────────────


def _is_rtl(target_language_code: str) -> bool:
    return target_language_code.lower() in RTL_LANGUAGE_CODES


def _mark_paragraph_rtl(paragraph: DocxParagraph) -> None:
    """Set `<w:bidi/>` on the paragraph's pPr so Word renders it right-to-left."""
    pPr = paragraph._p.get_or_add_pPr()  # noqa: SLF001 — required private API
    if pPr.find(qn("w:bidi")) is not None:
        return
    bidi = etree.SubElement(pPr, qn("w:bidi"))
    bidi.set(qn("w:val"), "1")


def _add_inline_runs(paragraph: DocxParagraph, text: str) -> None:
    """Add bold/italic-aware runs to a docx paragraph from inline markdown."""
    for run_spec in _parse_inline(text):
        run = paragraph.add_run(run_spec.text)
        run.bold = run_spec.bold
        run.italic = run_spec.italic


def _docx_style_for(kind: BlockKind, doc: Document) -> str:
    """Map a paragraph-shaped block kind to a python-docx built-in style name."""
    return {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "bullet": "List Bullet",
        "numbered": "List Number",
        "paragraph": "Normal",
    }[kind]


def _add_docx_table(doc: Document, block: Block, *, rtl: bool) -> None:
    """Render a markdown-table Block as a real docx Table."""
    if not block.rows:
        return
    cols = max(len(row) for row in block.rows)
    table = doc.add_table(rows=len(block.rows), cols=cols)
    # Built-in light grid style — visible borders, header band shading.
    try:
        table.style = doc.styles["Light Grid Accent 1"]
    except KeyError:  # pragma: no cover — older Word doesn't ship this style
        table.style = doc.styles["Table Grid"]

    for r_idx, row_data in enumerate(block.rows):
        row = table.rows[r_idx]
        is_header = r_idx == 0
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            # The cell starts with one empty paragraph; reuse it.
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            _add_inline_runs(paragraph, text)
            if is_header:
                for run in paragraph.runs:
                    run.bold = True
            if rtl:
                _mark_paragraph_rtl(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build_docx(title: str, body: str, target_language_code: str) -> bytes:
    """Render a .docx with a header block and structured body."""
    rtl = _is_rtl(target_language_code)
    doc = Document()

    heading = doc.add_heading(title, level=1)
    if rtl:
        _mark_paragraph_rtl(heading)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(_FOOTER_LINE)
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    if rtl:
        _mark_paragraph_rtl(subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")

    for block in _parse_blocks(body):
        if block.kind == "table":
            _add_docx_table(doc, block, rtl=rtl)
            continue
        style_name = _docx_style_for(block.kind, doc)
        # Headings use add_heading so Word's outline mode picks them up.
        if block.kind in ("h1", "h2"):
            level = 1 if block.kind == "h1" else 2
            para = doc.add_heading(level=level)
            _add_inline_runs(para, block.text)
        else:
            para = doc.add_paragraph(style=style_name)
            _add_inline_runs(para, block.text)
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(11)
        if rtl:
            _mark_paragraph_rtl(para)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────
# PDF rendering
# ──────────────────────────────────────────────────────────────────────────


def _pdf_escape(text: str) -> str:
    """Escape XML metacharacters reportlab's paraparser would interpret."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _build_pdf_table(
    block: Block,
    header_style: ParagraphStyle,
    body_style: ParagraphStyle,
) -> PdfTable:
    """Render a markdown-table Block as a reportlab Table flowable.

    Uses Paragraph cells (not raw strings) so inline bold/italic in cell
    text gets honored. Auto-computes column widths from the available
    page width so wide tables flow on Letter page-size.
    """
    if not block.rows:
        return PdfTable([[""]])  # never reached — _parse_blocks won't emit empty
    cols = max(len(row) for row in block.rows)
    available_width = LETTER[0] - 2 * inch
    col_width = available_width / max(cols, 1)
    data: list[list[Paragraph]] = []
    for r_idx, row in enumerate(block.rows):
        style = header_style if r_idx == 0 else body_style
        rendered_row = [
            Paragraph(_inline_to_xml(cell), style)
            for cell in row
        ]
        # Pad short rows so every line has the same column count.
        while len(rendered_row) < cols:
            rendered_row.append(Paragraph("", style))
        data.append(rendered_row)
    table = PdfTable(data, colWidths=[col_width] * cols, repeatRows=1)
    table.setStyle(
        TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#999999")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ])
    )
    return table


def _inline_to_xml(text: str) -> str:
    """Convert markdown inline (**bold**, *italic*) to reportlab `<b>`/`<i>` tags.

    We escape XML metacharacters FIRST so that body text with `<x>` doesn't
    crash the build, then inject `<b>` and `<i>` tags around the spans.
    Pattern intent: bold before italic so `**foo**` doesn't get eaten by
    the italic regex.
    """
    escaped = _pdf_escape(text)
    escaped = _BOLD_RE.sub(r"<b>\1</b>", escaped)
    escaped = _ITALIC_RE.sub(r"<i>\1</i>", escaped)
    return escaped


def build_pdf(title: str, body: str, target_language_code: str) -> bytes:
    """Render a .pdf with structured headings, lists, and inline styling."""
    rtl = _is_rtl(target_language_code)
    align = TA_RIGHT if rtl else TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=title,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TranslateTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        alignment=align,
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "TranslateSubtitle",
        parent=styles["Italic"],
        fontName="Helvetica-Oblique",
        fontSize=10,
        leading=12,
        alignment=align,
        spaceAfter=14,
    )
    h1_style = ParagraphStyle(
        "TranslateH1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        alignment=align,
        spaceBefore=10,
        spaceAfter=4,
    )
    h2_style = ParagraphStyle(
        "TranslateH2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        alignment=align,
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "TranslateBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        alignment=align,
        spaceAfter=8,
    )
    list_item_style = ParagraphStyle(
        "TranslateListItem",
        parent=body_style,
        spaceAfter=2,
    )

    flowables: list = [
        Paragraph(_pdf_escape(title), title_style),
        Paragraph(_pdf_escape(_FOOTER_LINE), subtitle_style),
        Spacer(1, 6),
    ]

    # Group consecutive list items so reportlab renders them as a single
    # ListFlowable with proper indentation/bullets, instead of one
    # awkward run-on per item.
    pending_bullets: list[Paragraph] = []
    pending_numbers: list[Paragraph] = []

    def flush_bullets() -> None:
        if pending_bullets:
            flowables.append(
                ListFlowable(
                    [ListItem(p, leftIndent=18) for p in pending_bullets],
                    bulletType="bullet",
                    leftIndent=18,
                )
            )
            pending_bullets.clear()

    def flush_numbers() -> None:
        if pending_numbers:
            flowables.append(
                ListFlowable(
                    [ListItem(p, leftIndent=22) for p in pending_numbers],
                    bulletType="1",
                    leftIndent=22,
                )
            )
            pending_numbers.clear()

    cell_style = ParagraphStyle(
        "TranslateCell",
        parent=body_style,
        spaceAfter=0,
        leading=13,
    )
    cell_header_style = ParagraphStyle(
        "TranslateCellHeader",
        parent=cell_style,
        fontName="Helvetica-Bold",
    )

    for block in _parse_blocks(body):
        if block.kind == "table":
            flush_bullets()
            flush_numbers()
            flowables.append(_build_pdf_table(block, cell_header_style, cell_style))
            continue
        rendered = _inline_to_xml(block.text)
        if block.kind == "bullet":
            flush_numbers()
            pending_bullets.append(Paragraph(rendered, list_item_style))
            continue
        if block.kind == "numbered":
            flush_bullets()
            pending_numbers.append(Paragraph(rendered, list_item_style))
            continue
        flush_bullets()
        flush_numbers()
        if block.kind == "h1":
            flowables.append(Paragraph(rendered, h1_style))
        elif block.kind == "h2":
            flowables.append(Paragraph(rendered, h2_style))
        else:
            flowables.append(Paragraph(rendered, body_style))

    flush_bullets()
    flush_numbers()

    doc.build(flowables)
    return buf.getvalue()
