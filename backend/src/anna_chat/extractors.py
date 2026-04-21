"""Text extraction dispatcher for uploaded attachments.

Each extractor returns plain text. The top-level `extract` function caps output
to `max_text_bytes` and signals whether truncation occurred.

Supported content types (see docs/ATTACHMENTS_CONTRACT.md):
  - application/pdf
  - image/png, image/jpeg
  - application/vnd.openxmlformats-officedocument.spreadsheetml.sheet  (xlsx)
  - application/vnd.openxmlformats-officedocument.wordprocessingml.document  (docx)
  - text/csv
  - text/plain
"""

from __future__ import annotations

import csv
import io
import zipfile
from collections.abc import Callable
from dataclasses import dataclass

import boto3
import defusedxml

# Harden stdlib XML parsers against XXE, billion-laughs, etc. python-docx and
# openpyxl delegate XML parsing to the stdlib / lxml; defusedxml monkey-patches
# the common entry points so any parse call in this process is safe-by-default.
defusedxml.defuse_stdlib()

DEFAULT_MAX_TEXT_BYTES = 500 * 1024

# Zip-bomb defenses for attacker-controlled OOXML (xlsx, docx). Both formats
# are ZIP containers; a tiny compressed payload can expand to gigabytes and
# OOM the 2 GB Lambda. We reject anything whose uncompressed total exceeds
# MAX_ZIP_TOTAL_BYTES or whose per-entry compression ratio exceeds
# MAX_ZIP_COMPRESSION_RATIO before handing the bytes to openpyxl / python-docx.
MAX_ZIP_TOTAL_BYTES = 250 * 1024 * 1024
MAX_ZIP_COMPRESSION_RATIO = 100

PDF_MIME = "application/pdf"
PNG_MIME = "image/png"
JPEG_MIME = "image/jpeg"
XLSX_MIME = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
CSV_MIME = "text/csv"
TXT_MIME = "text/plain"

ALLOWED_CONTENT_TYPES: frozenset[str] = frozenset(
    {PDF_MIME, PNG_MIME, JPEG_MIME, XLSX_MIME, DOCX_MIME, CSV_MIME, TXT_MIME}
)


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    truncated: bool


class UnsupportedContentType(Exception):
    pass


class ZipBombError(Exception):
    """Raised when a ZIP-based document (xlsx/docx) looks like a zip bomb.

    The message is intentionally generic so we don't echo attacker-crafted
    entry names or sizes back through logs or the UI.
    """

    pass


def _guard_office_zip(data: bytes) -> None:
    """Reject xlsx/docx payloads that look like zip bombs.

    Checks the zip central directory without extracting anything:
      - total uncompressed size across all entries <= MAX_ZIP_TOTAL_BYTES
      - per-entry compression ratio <= MAX_ZIP_COMPRESSION_RATIO

    Raises ZipBombError on violation.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            infos = zf.infolist()
    except zipfile.BadZipFile as exc:
        raise ZipBombError("document failed zip safety checks") from exc

    total = 0
    for info in infos:
        total += info.file_size
        if total > MAX_ZIP_TOTAL_BYTES:
            raise ZipBombError("document failed zip safety checks")
        ratio = info.file_size / max(info.compress_size, 1)
        if ratio > MAX_ZIP_COMPRESSION_RATIO:
            raise ZipBombError("document failed zip safety checks")


def extract(
    content_type: str,
    data: bytes,
    *,
    max_text_bytes: int = DEFAULT_MAX_TEXT_BYTES,
    textract_client=None,
) -> ExtractionResult:
    """Extract text from an uploaded attachment's bytes.

    Raises UnsupportedContentType for unknown MIME types. Callers should
    convert to a user-visible error.
    """
    handler = _DISPATCH.get(content_type)
    if handler is None:
        raise UnsupportedContentType(content_type)
    if content_type in {PDF_MIME, PNG_MIME, JPEG_MIME}:
        text = handler(data, textract_client=textract_client)
    else:
        text = handler(data)
    return _truncate(text, max_text_bytes)


def _truncate(text: str, max_bytes: int) -> ExtractionResult:
    encoded = text.encode("utf-8")
    if len(encoded) <= max_bytes:
        return ExtractionResult(text=text, truncated=False)
    # Truncate on byte boundary then drop any dangling partial codepoint.
    cut = encoded[:max_bytes]
    safe = cut.decode("utf-8", errors="ignore")
    return ExtractionResult(text=safe, truncated=True)


# ---------- per-type extractors ----------

def _extract_textract(data: bytes, *, textract_client=None) -> str:
    """Synchronous single-page OCR via Textract. Used for PNG/JPEG and as
    a last-ditch fallback for scanned PDFs whose pypdf pass yielded nothing.

    Note: Textract's synchronous DetectDocumentText API only accepts
    SINGLE-page PDFs. Multi-page PDFs must use the async API. We route
    multi-page PDFs through pypdf and only hit this path as a fallback
    for single-page image-based PDFs — anything larger that trips
    UnsupportedDocumentException here just means the PDF is a scanned
    multi-page doc and the caller will get an empty string back.
    """
    client = textract_client or boto3.client("textract")
    try:
        resp = client.detect_document_text(Document={"Bytes": data})
    except client.exceptions.UnsupportedDocumentException:
        # Scanned multi-page PDF — we can't handle this synchronously.
        # Returning "" lets the ingest pipeline decide what to do (it'll
        # mark the doc `error` with a helpful status message).
        return ""
    lines: list[str] = []
    for block in resp.get("Blocks", []):
        if block.get("BlockType") == "LINE":
            text = block.get("Text")
            if text:
                lines.append(text)
    return "\n".join(lines)


# Minimum pypdf output before we consider a PDF "text-based". Under this
# threshold we assume the PDF is image-only (scanned) and fall back to
# Textract for OCR. 120 chars ≈ one line of body text — enough to catch
# PDFs that only have running headers / page numbers extracted.
_PYPDF_MIN_CHARS = 120


def _extract_pdf(data: bytes, *, textract_client=None) -> str:
    """Text-first PDF extraction with a Textract fallback for scans.

    pypdf handles the 95%+ case: text-based PDFs produced by Word / LaTeX /
    Google Docs / academic publishers. It's pure Python, runs in the Lambda
    process (no network), and supports multi-page natively. If pypdf
    returns nothing (image-only scanned PDF), fall back to Textract.

    Intentionally broad exception handling — pypdf raises a zoo of errors
    on unusual / corrupt PDFs (TypeError, ValueError, KeyError, deep
    AssertionError from its tokenizer, etc.). We want one bad page /
    malformed trailer to yield "use Textract" rather than killing the
    whole ingest.
    """
    # Local import: pypdf is a heavy dep we don't want loading in the
    # critical path for chat. Only hit on ingest/extract Lambdas.
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        # is_encrypted can itself raise on PDFs with malformed trailers.
        if reader.is_encrypted:
            try:
                if reader.decrypt("") == 0:
                    return ""
            except Exception:  # pragma: no cover — encrypted paths are fiddly
                return ""

        parts: list[str] = []
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:
                # pypdf can raise on individual pages with weird fonts,
                # unusual encodings, or malformed content streams. Skip
                # that page rather than failing the whole doc.
                continue
            if text:
                parts.append(text)
        combined = "\n".join(parts).strip()
    except Exception:
        # PdfReader construction or a top-level property access failed.
        # Fall through to Textract — it's our last-ditch fallback for
        # PDFs pypdf can't parse at all.
        combined = ""

    if len(combined) >= _PYPDF_MIN_CHARS:
        return combined

    # Scanned / pypdf-unparseable PDF — try Textract OCR. Single-page
    # scans work; multi-page scans return "" (UnsupportedDocument).
    return _extract_textract(data, textract_client=textract_client)


def _extract_xlsx(data: bytes) -> str:
    # Local import: optional runtime dep, only when XLSX hits this path.
    from openpyxl import load_workbook

    _guard_office_zip(data)
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            # Skip completely empty rows.
            if any(cell != "" for cell in cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document  # python-docx

    _guard_office_zip(data)
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_csv(data: bytes) -> str:
    # Sniff decode; fall back to latin-1.
    try:
        decoded = data.decode("utf-8")
    except UnicodeDecodeError:
        decoded = data.decode("latin-1")
    buf = io.StringIO(decoded)
    reader = csv.reader(buf)
    lines: list[str] = []
    for row in reader:
        lines.append("\t".join(row))
    return "\n".join(lines)


def _extract_txt(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")


_DISPATCH: dict[str, Callable] = {
    PDF_MIME: _extract_pdf,
    PNG_MIME: _extract_textract,
    JPEG_MIME: _extract_textract,
    XLSX_MIME: _extract_xlsx,
    DOCX_MIME: _extract_docx,
    CSV_MIME: _extract_csv,
    TXT_MIME: _extract_txt,
}
