import io

import boto3
import pytest
from botocore.stub import Stubber

from anna_chat.extractors import (
    CSV_MIME,
    DOCX_MIME,
    JPEG_MIME,
    PDF_MIME,
    PNG_MIME,
    TXT_MIME,
    XLSX_MIME,
    ExtractionResult,
    UnsupportedContentType,
    extract,
)


def test_extract_csv_tab_separates_rows():
    data = b"a,b,c\n1,2,3\n4,5,6\n"
    result = extract(CSV_MIME, data)
    assert isinstance(result, ExtractionResult)
    assert result.truncated is False
    assert "a\tb\tc" in result.text
    assert "1\t2\t3" in result.text
    assert "4\t5\t6" in result.text


def test_extract_csv_handles_latin1_fallback():
    # Byte 0xE9 is é in latin-1 but invalid by itself as utf-8.
    data = b"name,note\nJos\xe9,hi\n"
    result = extract(CSV_MIME, data)
    assert "Jos" in result.text
    assert result.truncated is False


def test_extract_txt_utf8():
    data = b"hello world\nline two"
    result = extract(TXT_MIME, data)
    assert result.text == "hello world\nline two"
    assert result.truncated is False


def test_extract_txt_latin1_fallback():
    data = b"caf\xe9"  # 'café' in latin-1 only
    result = extract(TXT_MIME, data)
    assert "caf" in result.text
    assert result.truncated is False


def test_extract_txt_truncates_at_max_bytes():
    # ASCII so 1 char == 1 byte.
    data = ("x" * 1000).encode("utf-8")
    result = extract(TXT_MIME, data, max_text_bytes=100)
    assert result.truncated is True
    assert len(result.text.encode("utf-8")) <= 100


def test_extract_docx_reads_paragraphs_and_tables():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from a paragraph")
    doc.add_paragraph("Second paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "name"
    table.rows[0].cells[1].text = "value"
    table.rows[1].cells[0].text = "foo"
    table.rows[1].cells[1].text = "bar"
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    result = extract(DOCX_MIME, data)
    assert "Hello from a paragraph" in result.text
    assert "Second paragraph" in result.text
    assert "name\tvalue" in result.text
    assert "foo\tbar" in result.text
    assert result.truncated is False


def test_extract_xlsx_iterates_sheets_tab_separated():
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    ws1.title = "First"
    ws1.append(["h1", "h2"])
    ws1.append([1, 2])
    ws2 = wb.create_sheet("Second")
    ws2.append(["x", "y"])
    ws2.append(["a", "b"])
    buf = io.BytesIO()
    wb.save(buf)
    data = buf.getvalue()

    result = extract(XLSX_MIME, data)
    assert "# Sheet: First" in result.text
    assert "h1\th2" in result.text
    assert "1\t2" in result.text
    assert "# Sheet: Second" in result.text
    assert "a\tb" in result.text


def test_extract_pdf_uses_pypdf_for_text_based_pdfs(monkeypatch):
    # pypdf's PdfReader is imported inside _extract_pdf. Patch it there to
    # avoid needing a real PDF on disk — we're asserting the dispatch path,
    # not pypdf itself.
    class _FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    # Text must be >= 120 chars across all pages combined, otherwise the
    # extractor treats the PDF as "probably scanned" and falls back to
    # Textract. Make each page comfortably long.
    page_one_body = "Page one body " + ("lorem ipsum " * 10)
    page_two_body = "Page two body " + ("lorem ipsum " * 10)

    class _FakeReader:
        def __init__(self, _stream) -> None:
            self.is_encrypted = False
            self.pages = [_FakePage(page_one_body), _FakePage(page_two_body)]

    import pypdf

    monkeypatch.setattr(pypdf, "PdfReader", _FakeReader)
    result = extract(PDF_MIME, b"%PDF-1.4 fake")
    assert "Page one body" in result.text
    assert "Page two body" in result.text
    assert result.truncated is False


def test_extract_pdf_falls_back_to_textract_when_pypdf_empty(monkeypatch):
    # Scanned single-page PDF: pypdf yields < _PYPDF_MIN_CHARS, so we OCR
    # through Textract. Multi-page scans still return empty (sync Textract
    # caps at single-page), and that case is covered by the kb_ingest
    # NoTextExtracted error branch.
    class _FakePage:
        def extract_text(self) -> str:
            return ""  # image-only scan

    class _FakeReader:
        def __init__(self, _stream) -> None:
            self.is_encrypted = False
            self.pages = [_FakePage()]

    import pypdf

    monkeypatch.setattr(pypdf, "PdfReader", _FakeReader)

    client = boto3.client("textract", region_name="us-east-1")
    stub = Stubber(client)
    stub.add_response(
        "detect_document_text",
        {"Blocks": [{"BlockType": "LINE", "Text": "OCR'd line"}]},
        {"Document": {"Bytes": b"%PDF-1.4 fake"}},
    )
    stub.activate()
    try:
        result = extract(PDF_MIME, b"%PDF-1.4 fake", textract_client=client)
    finally:
        stub.deactivate()
    assert result.text == "OCR'd line"


def test_extract_png_uses_textract_via_stub():
    client = boto3.client("textract", region_name="us-east-1")
    stub = Stubber(client)
    stub.add_response(
        "detect_document_text",
        {"Blocks": [{"BlockType": "LINE", "Text": "png text"}]},
        {"Document": {"Bytes": b"\x89PNG..."}},
    )
    stub.activate()
    try:
        result = extract(PNG_MIME, b"\x89PNG...", textract_client=client)
    finally:
        stub.deactivate()
    assert result.text == "png text"


def test_extract_jpeg_uses_textract_via_stub():
    client = boto3.client("textract", region_name="us-east-1")
    stub = Stubber(client)
    stub.add_response(
        "detect_document_text",
        {"Blocks": [{"BlockType": "LINE", "Text": "jpeg text"}]},
        {"Document": {"Bytes": b"\xff\xd8\xff..."}},
    )
    stub.activate()
    try:
        result = extract(JPEG_MIME, b"\xff\xd8\xff...", textract_client=client)
    finally:
        stub.deactivate()
    assert result.text == "jpeg text"


def test_extract_rejects_unsupported_content_type():
    with pytest.raises(UnsupportedContentType):
        extract("application/x-rar", b"\x00")


def test_extract_truncation_flag_on_textract_output():
    # PNG → Textract sync (single image), so this exercises the same
    # truncation wrapper as before without needing a multi-page PDF stub.
    client = boto3.client("textract", region_name="us-east-1")
    stub = Stubber(client)
    long_line = "A" * 200
    blocks = [{"BlockType": "LINE", "Text": long_line} for _ in range(20)]
    stub.add_response(
        "detect_document_text",
        {"Blocks": blocks},
        {"Document": {"Bytes": b"\x89PNG..."}},
    )
    stub.activate()
    try:
        result = extract(
            PNG_MIME, b"\x89PNG...", textract_client=client, max_text_bytes=500
        )
    finally:
        stub.deactivate()
    assert result.truncated is True
    assert len(result.text.encode("utf-8")) <= 500
