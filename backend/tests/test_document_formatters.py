import io

import pytest
from docx import Document
from docx.oxml.ns import qn

from anna_chat.document_formatters import (
    RTL_LANGUAGE_CODES,
    _parse_blocks,
    _parse_inline,
    build_docx,
    build_pdf,
)


def _has_bidi(paragraph) -> bool:
    """True if the paragraph has a `<w:bidi/>` element in its pPr (RTL marker)."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:bidi")) is not None


# ---------- block parser ----------


def test_parse_blocks_blank_line_boundaries():
    body = "First para line one.\nFirst para line two.\n\nSecond para.\n\n\nThird para."
    blocks = _parse_blocks(body)
    assert len(blocks) == 3
    assert all(b.kind == "paragraph" for b in blocks)
    # Lines within a paragraph are joined with a space.
    assert blocks[0].text.startswith("First para line one")
    assert "First para line two" in blocks[0].text
    assert blocks[1].text == "Second para."
    assert blocks[2].text == "Third para."


def test_parse_blocks_empty_input_returns_empty_list():
    assert _parse_blocks("") == []
    assert _parse_blocks("   \n\n   ") == []


def test_parse_blocks_h1_and_h2_headings():
    body = "# Background\n\n## Subsection\n\nBody paragraph."
    blocks = _parse_blocks(body)
    assert [b.kind for b in blocks] == ["h1", "h2", "paragraph"]
    assert blocks[0].text == "Background"
    assert blocks[1].text == "Subsection"


def test_parse_blocks_bullet_list():
    body = "Lead paragraph.\n\n- First item\n- Second item\n- Third item\n\nClosing."
    blocks = _parse_blocks(body)
    kinds = [b.kind for b in blocks]
    assert kinds == ["paragraph", "bullet", "bullet", "bullet", "paragraph"]
    assert blocks[1].text == "First item"
    assert blocks[3].text == "Third item"


def test_parse_blocks_numbered_list():
    body = "1. First step\n2. Second step\n10. Tenth step"
    blocks = _parse_blocks(body)
    assert [b.kind for b in blocks] == ["numbered", "numbered", "numbered"]
    assert blocks[0].text == "First step"
    assert blocks[2].text == "Tenth step"


def test_parse_blocks_mixed_structure():
    body = (
        "# Title\n\n"
        "Intro paragraph.\n\n"
        "## Methods\n\n"
        "- Point one\n- Point two\n\n"
        "Closing thought."
    )
    blocks = _parse_blocks(body)
    assert [b.kind for b in blocks] == [
        "h1", "paragraph", "h2", "bullet", "bullet", "paragraph"
    ]


# ---------- inline parser ----------


def test_parse_inline_plain_text():
    runs = _parse_inline("hello world")
    assert len(runs) == 1
    assert runs[0].text == "hello world"
    assert runs[0].bold is False
    assert runs[0].italic is False


def test_parse_inline_bold():
    runs = _parse_inline("a **bold** word")
    texts = [(r.text, r.bold, r.italic) for r in runs]
    assert texts == [("a ", False, False), ("bold", True, False), (" word", False, False)]


def test_parse_inline_italic():
    runs = _parse_inline("see *Crank et al* paper")
    texts = [(r.text, r.bold, r.italic) for r in runs]
    assert texts == [
        ("see ", False, False),
        ("Crank et al", False, True),
        (" paper", False, False),
    ]


def test_parse_inline_mixed_bold_and_italic():
    runs = _parse_inline("**Term**: see *paper title* for details")
    # Two style spans + plain in between
    assert any(r.bold and r.text == "Term" for r in runs)
    assert any(r.italic and r.text == "paper title" for r in runs)


def test_parse_inline_bold_does_not_eat_italic():
    """Regression: `**foo**` must not be parsed as `*` + `*foo*` + `*`."""
    runs = _parse_inline("**foo**")
    assert len(runs) == 1
    assert runs[0].text == "foo"
    assert runs[0].bold is True
    assert runs[0].italic is False


# ---------- docx ----------


def test_build_docx_returns_valid_docx_with_title_and_body():
    title = "Crank et al 2021.pdf (Spanish)"
    body = "Hola.\n\nEsto es un párrafo.\n\nÚltimo párrafo."
    raw = build_docx(title, body, target_language_code="es")
    assert isinstance(raw, bytes)
    assert len(raw) > 0
    doc = Document(io.BytesIO(raw))
    rendered = "\n".join(p.text for p in doc.paragraphs)
    assert title in rendered
    assert "Translated by Praxis" in rendered
    assert "Hola." in rendered
    assert "Último párrafo." in rendered


def test_build_docx_renders_headings_with_heading_style():
    body = "# Section A\n\nBody under A.\n\n## Subsection\n\nMore body."
    raw = build_docx("Doc", body, target_language_code="es")
    doc = Document(io.BytesIO(raw))
    # python-docx assigns built-in heading paragraphs the style name
    # "Heading 1" / "Heading 2". The doc title also uses Heading 1, so
    # we look for at least two H1 paragraphs (title + Section A) and
    # at least one H2.
    h1_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 1")
    h2_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 2")
    assert h1_count >= 2
    assert h2_count >= 1


def test_build_docx_renders_bullets_with_list_style():
    body = "- Apple\n- Banana\n- Cherry"
    raw = build_docx("Doc", body, target_language_code="es")
    doc = Document(io.BytesIO(raw))
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullet_paras) == 3


def test_build_docx_renders_inline_bold():
    body = "This has a **bolded** word."
    raw = build_docx("Doc", body, target_language_code="es")
    doc = Document(io.BytesIO(raw))
    bold_runs = [
        run
        for p in doc.paragraphs
        for run in p.runs
        if run.bold and run.text == "bolded"
    ]
    assert len(bold_runs) == 1


def test_build_docx_rtl_target_sets_rtl_paragraph_format():
    raw = build_docx(
        "Doc title (Arabic)",
        "اختبار.\n\nجملة ثانية.",
        target_language_code="ar",
    )
    doc = Document(io.BytesIO(raw))
    assert any(_has_bidi(p) for p in doc.paragraphs), (
        "expected at least one paragraph with <w:bidi/> set for ar target"
    )


def test_build_docx_non_rtl_target_does_not_set_rtl():
    raw = build_docx(
        "Doc title (Spanish)",
        "Hola.\n\nMundo.",
        target_language_code="es",
    )
    doc = Document(io.BytesIO(raw))
    assert not any(_has_bidi(p) for p in doc.paragraphs)


# ---------- pdf ----------


def test_build_pdf_returns_valid_pdf_bytes():
    title = "Crank et al 2021.pdf (Spanish)"
    body = "Hola.\n\nEsto es un párrafo.\n\nÚltimo párrafo."
    raw = build_pdf(title, body, target_language_code="es")
    assert isinstance(raw, bytes)
    assert len(raw) > 0
    assert raw[:5] == b"%PDF-"
    assert b"%%EOF" in raw[-1024:]


def test_build_pdf_handles_empty_body():
    raw = build_pdf("Empty (Spanish)", "", target_language_code="es")
    assert raw[:5] == b"%PDF-"


def test_build_pdf_escapes_html_like_tokens_in_body():
    body = "Notes: <patient> & <provider> discussed > 3 options."
    raw = build_pdf("Notes (Spanish)", body, target_language_code="es")
    assert raw[:5] == b"%PDF-"


def test_build_pdf_renders_full_markdown_structure():
    """End-to-end PDF build with headings, lists, and inline emphasis —
    if anything regresses on parsing/escaping, reportlab will throw."""
    body = (
        "# Background\n\n"
        "Per **the contract**, treatment was given.\n\n"
        "## Methods\n\n"
        "- First method\n- Second method with *emphasis*\n\n"
        "1. Step one\n2. Step two\n\n"
        "Closing paragraph."
    )
    raw = build_pdf("Mixed (Spanish)", body, target_language_code="es")
    assert raw[:5] == b"%PDF-"
    assert b"%%EOF" in raw[-1024:]


def test_arabic_in_rtl_set():
    assert "ar" in RTL_LANGUAGE_CODES


@pytest.mark.parametrize("lang", ["es", "zh", "vi", "fr", "ja", "en"])
def test_build_docx_supports_multiple_target_languages(lang: str):
    raw = build_docx(
        f"Doc ({lang})",
        "Body content for the document.",
        target_language_code=lang,
    )
    doc = Document(io.BytesIO(raw))
    rendered = "\n".join(p.text for p in doc.paragraphs)
    assert "Body content for the document." in rendered
